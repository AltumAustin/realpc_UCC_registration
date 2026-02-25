"""Generate a well-formatted .docx describing the UCC data ingestion approach."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "UCC_Data_Ingestion_Approach.docx")


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_color="1B3A5C"):
    """Style a table header row with dark background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def style_data_row(row, alt=False):
    """Style a data row with optional alternating background."""
    bg = "F2F6FA" if alt else "FFFFFF"
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
    style_header_row(header_row)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
        style_data_row(row, alt=(row_idx % 2 == 1))

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph("")  # spacer
    return table


def add_key_value_line(doc, key, value):
    """Add a bold key: normal value line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(10)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(10)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def build_document():
    doc = Document()

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        hs.font.name = "Calibri"

    # ======================================================================
    # TITLE PAGE
    # ======================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UCC Filing Data\nIngestion Strategy")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Recommended Approach for Nationwide Coverage")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x6B, 0x82)

    doc.add_paragraph("")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared for Real Finance, Inc.\nFebruary 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x6B, 0x82)

    doc.add_paragraph("")
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("CONFIDENTIAL")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.bold = True

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS (manual)
    # ======================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  Executive Summary",
        "2.  Objective & Constraints",
        "3.  Recommended Approach: Hybrid 3-Tier Architecture",
        "4.  Tier 1 — Free Open Data APIs",
        "5.  Tier 2 — State Bulk Subscriptions",
        "6.  Tier 3 — Commercial Data Provider",
        "7.  Commercial Provider Comparison",
        "8.  Data Normalization & Storage",
        "9.  Cost Summary",
        "10. Expected Latency by Source",
        "11. Rollout Plan",
        "12. Risk Factors & Mitigations",
        "13. Recommendation",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ======================================================================
    # 1. EXECUTIVE SUMMARY
    # ======================================================================
    doc.add_heading("1. Executive Summary", level=1)

    p = doc.add_paragraph(
        "Real Finance, Inc. requires continuous access to every UCC (Uniform Commercial Code) "
        "filing across all 50 U.S. states and the District of Columbia. This document outlines "
        "the recommended approach for building a centralized UCC filing database that captures "
        "new filings within three days of their recording, at an estimated annual cost of "
        "approximately $399,000 — well within the $1,000,000 budget."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph(
        "The recommended architecture is a hybrid 3-tier ingestion system that strategically "
        "combines free government APIs, direct state bulk data subscriptions, and a commercial "
        "data provider to optimize for cost, coverage, and latency. This approach avoids "
        "over-reliance on any single data source while ensuring that all 51 jurisdictions are "
        "covered with no gaps."
    )

    # ======================================================================
    # 2. OBJECTIVE & CONSTRAINTS
    # ======================================================================
    doc.add_heading("2. Objective & Constraints", level=1)

    doc.add_heading("Objective", level=2)
    p = doc.add_paragraph(
        "Continuously ingest every UCC filing from all 51 jurisdictions (50 states + DC) "
        "into a centralized, normalized database."
    )

    doc.add_heading("Constraints", level=2)
    add_bullet(doc, "$1,000,000 per year", bold_prefix="Maximum annual budget: ")
    add_bullet(doc, "3 calendar days from date of filing at the state level", bold_prefix="Maximum acceptable latency: ")
    add_bullet(doc, "~145,000 new filings per month (~1.7 million per year)", bold_prefix="Expected data volume: ")
    add_bullet(doc, "All records must be normalized to a common schema regardless of source format", bold_prefix="Data quality: ")
    add_bullet(doc, "Full audit trail of every ingestion run, with original source records preserved", bold_prefix="Auditability: ")

    # ======================================================================
    # 3. RECOMMENDED APPROACH
    # ======================================================================
    doc.add_heading("3. Recommended Approach: Hybrid 3-Tier Architecture", level=1)

    p = doc.add_paragraph(
        "Rather than relying on a single commercial provider for all 51 jurisdictions — which "
        "would cost $300K–$500K annually and create a single point of failure — the recommended "
        "approach uses three tiers of data sources, prioritizing the cheapest and most direct "
        "sources wherever they are available."
    )
    doc.add_paragraph("")

    add_styled_table(
        doc,
        ["Tier", "Source Type", "States", "Annual Cost", "Latency"],
        [
            ["Tier 1", "Free Open Data APIs (Socrata)", "2", "$0", "~1 day"],
            ["Tier 2", "State Bulk Subscriptions", "15", "~$98,830", "1–3 days"],
            ["Tier 3", "Commercial Provider API", "~34", "~$300,000", "1–3 days"],
            ["Total", "All sources combined", "51", "~$398,830", "≤3 days"],
        ],
        col_widths=[2.0, 5.5, 2.0, 3.0, 2.5],
    )

    p = doc.add_paragraph(
        "This tiered approach provides a ~$600,000 budget buffer to absorb price increases, "
        "higher-than-expected commercial provider costs, and future infrastructure expenses."
    )

    doc.add_heading("Why a Hybrid Approach?", level=2)
    add_bullet(doc, "Free APIs for CT and CO eliminate cost for those states entirely and provide the freshest data (daily updates).", bold_prefix="Cost optimization: ")
    add_bullet(doc, "15 states offer bulk data subscriptions at far lower cost than commercial providers. For example, California's weekly XML downloads are free, and North Dakota is only $480/year.", bold_prefix="Direct state relationships: ")
    add_bullet(doc, "Where Tier 2 and Tier 3 overlap, we can cross-reference records to identify data quality issues from either source.", bold_prefix="Data quality cross-checks: ")
    add_bullet(doc, "No single vendor can cut off access to all 51 jurisdictions. Free APIs and state bulk subscriptions are not subject to commercial contract disputes.", bold_prefix="Reduced vendor risk: ")

    # ======================================================================
    # 4. TIER 1 — FREE OPEN DATA APIs
    # ======================================================================
    doc.add_heading("4. Tier 1 — Free Open Data APIs", level=1)

    p = doc.add_paragraph(
        "Two states — Connecticut and Colorado — publish their UCC filing indexes as open "
        "data via the Socrata SODA API platform. These are free, require no authentication "
        "(though an optional app token raises rate limits), and provide daily-refreshed JSON data."
    )

    add_styled_table(
        doc,
        ["State", "API Platform", "Frequency", "Latency", "Format", "Annual Cost"],
        [
            ["Connecticut", "data.ct.gov (Socrata SODA)", "Daily", "~1 day", "JSON", "$0"],
            ["Colorado", "data.colorado.gov (Socrata SODA)", "Daily", "~1.5 days", "JSON", "$0"],
        ],
        col_widths=[2.5, 5.0, 2.0, 2.0, 1.5, 2.0],
    )

    doc.add_heading("How It Works", level=2)
    add_bullet(doc, "The ingestion scheduler polls each API daily using the Socrata SODA query language.")
    add_bullet(doc, "Incremental pulls are performed using a date filter ($where clause) to fetch only filings recorded since the last ingestion run.")
    add_bullet(doc, "Pagination is handled via $limit and $offset parameters (default page size: 5,000 records).")
    add_bullet(doc, "No credentials are required. An optional Socrata app token raises the rate limit from 1,000 to 4,000 requests per hour.")

    p = doc.add_paragraph()
    run = p.add_run("Total Tier 1 cost: $0/year")
    run.bold = True

    # ======================================================================
    # 5. TIER 2 — STATE BULK SUBSCRIPTIONS
    # ======================================================================
    doc.add_heading("5. Tier 2 — State Bulk Subscriptions", level=1)

    p = doc.add_paragraph(
        "Fifteen states offer direct bulk data access — either as downloadable files (CSV, XML, "
        "fixed-width, tab-delimited) or via subscription-based data feeds. These are purchased "
        "directly from the state Secretary of State office (or equivalent agency) and provide "
        "the most authoritative data available."
    )

    doc.add_heading("Data Delivery Models", level=2)
    p = doc.add_paragraph(
        "State bulk data programs use one of three delivery models. Understanding which model "
        "each state uses is critical for efficient ingestion — it determines whether we need "
        "an initial data purchase, how much data each download contains, and how the pipeline "
        "processes it."
    )

    add_styled_table(
        doc,
        ["Delivery Model", "States", "How It Works", "Efficiency"],
        [
            [
                "Initial Load +\nIncremental Updates",
                "TX, CA, MN, FL",
                "Purchase a one-time master file to seed the\ndatabase, then subscribe to periodic updates\ncontaining only new/changed filings.",
                "Most efficient — daily\ndownloads are small\nand fast to process.",
            ],
            [
                "Full Replacement",
                "ID",
                "Each extract is the COMPLETE database.\nNo initial load needed, but every download\nreplaces all prior data.",
                "Simple but less efficient\n— re-parses all records\neach cycle.",
            ],
            [
                "To Be Verified",
                "KY, WV, ND, AR,\nIN, NY, NC, SC,\nSD, AZ",
                "These states offer bulk subscriptions but we\nhave not yet confirmed whether they deliver\nfull dumps or incremental deltas.",
                "Will be verified during\nsubscription setup for\neach state.",
            ],
        ],
        col_widths=[3.0, 3.0, 5.5, 3.5],
    )

    p = doc.add_paragraph(
        "The ingestion pipeline handles all three models safely. Every record is stored via "
        "an upsert operation (insert if new, update if existing) keyed on the unique combination "
        "of filing number, state, debtor name, and secured party name. This means that even "
        "full-replacement downloads will not create duplicate records — existing records are "
        "simply updated in place."
    )

    doc.add_heading("Per-State Details", level=2)

    add_styled_table(
        doc,
        ["State", "Format", "Frequency", "Latency", "Annual Cost", "Notes"],
        [
            ["California", "XML", "Weekly", "≤3 days", "$0", "Master file ($100) + free weekly updates"],
            ["Texas", "JSON", "Daily", "1 day", "$1,350", "Master Unload + Daily Filing Data Updates"],
            ["Kentucky", "CSV", "Daily", "1 day", "$18,000", "$1,500/mo; delivery model TBD"],
            ["West Virginia", "CSV", "Weekly", "≤3 days", "$5,000", "3 files per week; delivery model TBD"],
            ["Idaho", "Tab-delimited", "Biweekly", "≤3 days", "$3,250", "Full replacement each extract"],
            ["North Dakota", "CSV", "Biweekly", "≤3 days", "$480", "$40/mo; delivery model TBD"],
            ["Minnesota", "CSV", "Weekly", "≤3 days", "$5,000", "Initial dataset + weekly updates"],
            ["Arkansas", "CSV", "Weekly", "≤3 days", "$150", "$150/yr INA; delivery model TBD"],
            ["Indiana", "XML (IACA v4.0)", "Weekly", "≤3 days", "$2,000", "REST API via INBiz; model TBD"],
            ["New York", "XML", "Weekly", "≤3 days", "$3,600", "$300/mo subscription; model TBD"],
            ["North Carolina", "CSV", "Weekly", "≤3 days", "$5,000", "Contract-based; model TBD"],
            ["South Carolina", "CSV", "Weekly", "≤3 days", "$54,000", "$4,500/mo; delivery model TBD"],
            ["South Dakota", "CSV", "Weekly", "≤3 days", "$1,000", "Subscription-based; model TBD"],
            ["Arizona", "CSV", "Monthly", "≤3 days*", "$2,000", "*Monthly cadence; model TBD"],
            ["Florida", "Fixed-width", "Daily", "1 day", "$3,000", "Quarterly full dump + daily updates"],
        ],
        col_widths=[2.3, 2.2, 1.8, 1.5, 2.0, 5.2],
    )

    doc.add_heading("Subscription Priority Order", level=2)
    p = doc.add_paragraph(
        "We recommend activating Tier 2 subscriptions in the following order, based on cost "
        "efficiency and data freshness:"
    )
    add_bullet(doc, "California (free), Texas ($1,350 Master Unload + Daily Filing Data Updates) — highest value, daily/weekly data", bold_prefix="Priority 1: ")
    add_bullet(doc, "Kentucky ($18K), Florida ($3K) — daily data, important jurisdictions", bold_prefix="Priority 2: ")
    add_bullet(doc, "New York ($3,600) — large filing volume state", bold_prefix="Priority 3: ")
    add_bullet(doc, "WV, MN, NC, IN, AR, ND, SD, ID (~$21K combined)", bold_prefix="Priority 4: ")
    add_bullet(doc, "South Carolina ($54K) — evaluate if commercial provider covers adequately first", bold_prefix="Priority 5: ")
    add_bullet(doc, "Arizona ($2K) — monthly cadence may require commercial backup", bold_prefix="Priority 6: ")

    doc.add_heading("Texas: Master Unload + Daily Filing Data Updates", level=2)
    p = doc.add_paragraph(
        "Texas uses a two-step process managed by the Secretary of State's office. "
        "This is the model for how we establish and maintain the Texas portion of "
        "our UCC database:"
    )

    add_bullet(
        doc,
        "Purchase the Master Unload ($1,350) from the Texas SOS. This is a one-time "
        "purchase containing every historical UCC filing in the SOS database. It seeds "
        "our local database with the complete Texas filing index.",
        bold_prefix="Step 1 — Master Unload: ",
    )
    add_bullet(
        doc,
        "Subscribe to the Daily Filing Data Updates. This is an ongoing subscription "
        "that delivers a daily incremental JSON file containing only filings that were "
        "recorded, amended, or updated since the previous day. These daily updates are "
        "downloaded and applied against the Master Unload to keep our local database "
        "current with the data maintained in the SOS database.",
        bold_prefix="Step 2 — Daily Filing Data Updates: ",
    )

    doc.add_paragraph("")
    p = doc.add_paragraph(
        "Both the Master Unload and Daily Filing Data Updates use the same JSON schema, "
        "so the ingestion pipeline handles them with a single parser. The pipeline "
        "automatically detects which mode to use: if no Texas data exists in the local "
        "database, it expects a Master Unload file; if Texas data already exists, it "
        "looks for the latest Daily Filing Data Update."
    )

    p = doc.add_paragraph()
    run = p.add_run("Total Tier 2 cost: ~$98,830/year")
    run.bold = True

    # ======================================================================
    # 6. TIER 3 — COMMERCIAL DATA PROVIDER
    # ======================================================================
    doc.add_heading("6. Tier 3 — Commercial Data Provider", level=1)

    p = doc.add_paragraph(
        "The remaining ~34 states do not offer practical direct bulk data access. For these "
        "jurisdictions, we recommend contracting with a single commercial UCC data provider "
        "that aggregates filings nationwide and exposes them via a REST/JSON API."
    )

    doc.add_paragraph("")

    p = doc.add_paragraph(
        "The commercial provider acts as a catch-all — it covers every state that we don't "
        "have a cheaper, more direct source for. As we activate more Tier 2 subscriptions over "
        "time, the commercial provider's role (and negotiating leverage for a lower contract price) "
        "can be adjusted."
    )

    doc.add_heading("States Covered by Tier 3", level=2)
    tier3_states = [
        "Alabama", "Alaska", "Delaware", "DC", "Georgia", "Hawaii", "Iowa",
        "Illinois", "Kansas", "Louisiana", "Maine", "Maryland", "Massachusetts",
        "Michigan", "Mississippi", "Missouri", "Montana", "Nebraska", "Nevada",
        "New Hampshire", "New Jersey", "New Mexico", "Ohio", "Oklahoma", "Oregon",
        "Pennsylvania", "Rhode Island", "Tennessee", "Utah", "Vermont", "Virginia",
        "Washington", "Wisconsin", "Wyoming"
    ]
    p = doc.add_paragraph(", ".join(tier3_states))
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x6B, 0x82)

    # ======================================================================
    # 7. COMMERCIAL PROVIDER COMPARISON
    # ======================================================================
    doc.add_heading("7. Commercial Provider Comparison", level=1)

    p = doc.add_paragraph(
        "Three commercial providers have been identified as suitable candidates for the "
        "Tier 3 data feed. We recommend soliciting quotes from all three and selecting based "
        "on price, API quality, and data freshness."
    )

    add_styled_table(
        doc,
        ["Provider", "Coverage", "API Style", "Est. Annual Cost", "Key Strengths"],
        [
            ["Baselayer", "50 states", "REST / JSON", "$200K–$400K", "API-first, modern platform, daily refresh"],
            ["First Corporate\nSolutions (FCS)", "50 states", "REST / JSON", "$200K–$400K", "Purchases direct from states; also covers Delaware"],
            ["LexisNexis", "48 states daily", "REST / JSON", "$300K–$500K", "Deepest historical archive, most comprehensive"],
        ],
        col_widths=[2.8, 2.0, 2.0, 3.0, 5.2],
    )

    doc.add_heading("Recommendation", level=2)
    p = doc.add_paragraph(
        "Begin with Baselayer or First Corporate Solutions due to their lower price range "
        "and modern API design. FCS has the additional advantage of being an existing Delaware "
        "Authorized Searcher, which is relevant because Delaware is the only state that restricts "
        "UCC search access to a closed set of 25 certified searchers — FCS is one of them."
    )

    # ======================================================================
    # 8. DATA NORMALIZATION & STORAGE
    # ======================================================================
    doc.add_heading("8. Data Normalization & Storage", level=1)

    p = doc.add_paragraph(
        "Data arrives in at least five different formats (JSON, CSV, XML, tab-delimited, "
        "fixed-width ASCII) with inconsistent field names, date formats, and status codes. "
        "A normalization layer standardizes all records before storage."
    )

    doc.add_heading("Normalization Rules", level=2)
    add_bullet(doc, "All dates converted to ISO 8601 format (YYYY-MM-DD)", bold_prefix="Dates: ")
    add_bullet(doc, "Mapped to standard codes: UCC-1, UCC-3, UCC-3 Amendment, UCC-3 Assignment, UCC-3 Continuation, UCC-3 Termination, UCC-5", bold_prefix="Filing types: ")
    add_bullet(doc, "Normalized to: active, lapsed, terminated, continued", bold_prefix="Status codes: ")
    add_bullet(doc, "Whitespace normalized; extra spaces collapsed", bold_prefix="Names: ")
    add_bullet(doc, "Standardized to 5-digit or ZIP+4 format", bold_prefix="ZIP codes: ")

    doc.add_heading("Database Schema", level=2)
    p = doc.add_paragraph(
        "All records are stored in a SQLite database with the following normalized schema:"
    )

    add_styled_table(
        doc,
        ["Field Group", "Fields", "Purpose"],
        [
            ["Identity", "filing_number, state", "Unique filing identification"],
            ["Filing Details", "filing_type, filing_date, lapse_date, filing_status", "Core filing metadata"],
            ["Debtor Info", "debtor_name, debtor_address, debtor_city,\ndebtor_state, debtor_zip, debtor_type", "Who owes the debt"],
            ["Secured Party", "secured_party_name, secured_party_address,\nsecured_party_city, secured_party_state, secured_party_zip", "Who holds the lien"],
            ["Collateral", "collateral_description", "What secures the filing"],
            ["Amendments", "original_filing_number, amendment_type", "Links to parent filing"],
            ["Audit", "source_tier, source_raw, ingested_at, last_updated_at", "Full traceability to original source"],
        ],
        col_widths=[2.5, 5.5, 5.0],
    )

    p = doc.add_paragraph(
        "The source_raw column preserves the original record exactly as received from the "
        "data source, ensuring that any normalization decisions can be audited or reversed."
    )

    doc.add_heading("Storage Capacity", level=2)
    add_bullet(doc, "~145,000 new records per month (~1.7 million per year)", bold_prefix="Expected volume: ")
    add_bullet(doc, "SQLite with WAL (Write-Ahead Logging) mode for concurrent read access", bold_prefix="Initial storage: ")
    add_bullet(doc, "PostgreSQL, if concurrent write access or higher volumes demand it", bold_prefix="Migration path: ")

    # ======================================================================
    # 9. COST SUMMARY
    # ======================================================================
    doc.add_heading("9. Cost Summary", level=1)

    add_styled_table(
        doc,
        ["Tier", "Source", "States", "Annual Cost"],
        [
            ["Tier 1", "Free Open Data APIs (Socrata)", "2", "$0"],
            ["Tier 2", "State Bulk Subscriptions", "15", "$98,830"],
            ["Tier 3", "Commercial Provider (est.)", "~34", "$300,000"],
            ["", "", "Total", "$398,830"],
            ["", "", "Budget ($1M) Remaining", "$601,170"],
        ],
        col_widths=[2.0, 5.5, 3.0, 3.0],
    )

    doc.add_heading("Budget Buffer Allocation", level=2)
    p = doc.add_paragraph(
        "The ~$600K buffer provides substantial headroom for:"
    )
    add_bullet(doc, "Higher-than-estimated commercial provider contract costs (quotes range $200K–$500K)")
    add_bullet(doc, "Additional Tier 2 subscriptions as states open new bulk access programs")
    add_bullet(doc, "Infrastructure and compute costs (servers, storage, monitoring)")
    add_bullet(doc, "Annual price increases from state agencies and commercial providers")
    add_bullet(doc, "Potential migration to managed PostgreSQL hosting")

    # ======================================================================
    # 10. EXPECTED LATENCY
    # ======================================================================
    doc.add_heading("10. Expected Latency by Source", level=1)

    p = doc.add_paragraph(
        "All 51 jurisdictions meet the ≤3 calendar day latency target:"
    )

    add_styled_table(
        doc,
        ["Data Source", "States", "Poll Frequency", "Typical Latency"],
        [
            ["Socrata API", "CT, CO", "Every 6 hours", "~1 day"],
            ["State daily downloads", "TX, KY, FL", "Daily (2 AM ET)", "1 day"],
            ["State weekly downloads", "CA, WV, MN, AR, IN, NY, NC, SC, SD", "Daily check, weekly data", "1–3 days"],
            ["State biweekly downloads", "ID, ND", "Daily check, biweekly data", "1–3 days"],
            ["Commercial provider API", "~34 states", "Daily (2 AM ET)", "1–3 days"],
        ],
        col_widths=[3.5, 5.0, 3.0, 2.5],
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Note: Arizona's monthly state subscription cadence may intermittently exceed "
        "3 days. The commercial provider serves as a backup source for Arizona to ensure "
        "the latency SLA is met."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x6B, 0x82)

    # ======================================================================
    # 11. ROLLOUT PLAN
    # ======================================================================
    doc.add_heading("11. Rollout Plan", level=1)

    p = doc.add_paragraph(
        "The rollout is structured in five phases over approximately eight weeks, with the "
        "free and lowest-risk sources activated first:"
    )

    doc.add_heading("Phase 1 — Validation (Week 1)", level=2)
    add_bullet(doc, "Test ingestion pipeline against Connecticut and Colorado free APIs")
    add_bullet(doc, "Validate database schema and normalization with real filing data")
    add_bullet(doc, "Verify incremental pull logic (fetch only new records since last run)")
    add_bullet(doc, "Cost: $0")

    doc.add_heading("Phase 2 — Open APIs Live (Week 2)", level=2)
    add_bullet(doc, "Set up automated cron schedule for CT and CO (every 6 hours)")
    add_bullet(doc, "Register for Socrata app token to raise API rate limits")
    add_bullet(doc, "Monitor ingestion logs for errors and data quality issues")
    add_bullet(doc, "Cost: $0")

    doc.add_heading("Phase 3 — Commercial Provider (Weeks 3–6)", level=2)
    add_bullet(doc, "Contact Baselayer, First Corporate Solutions, and LexisNexis for formal quotes")
    add_bullet(doc, "Negotiate and sign contract for nationwide daily data feed")
    add_bullet(doc, "Configure API credentials and run test connection")
    add_bullet(doc, "Execute initial full pull for all ~34 Tier 3 states")
    add_bullet(doc, "Set up daily automated ingestion cron job")
    add_bullet(doc, "Cost: ~$300,000/year (contract-dependent)")

    doc.add_heading("Phase 4 — State Subscriptions (Weeks 4–8)", level=2)
    add_bullet(doc, "Activate Tier 2 subscriptions in priority order (see Section 5)")
    add_bullet(doc, "Begin with California (free) and Texas ($1,350), then expand")
    add_bullet(doc, "Verify each state's delivery model (initial load + incremental vs. full replacement) during subscription setup")
    add_bullet(doc, "Cross-reference Tier 2 direct data against Tier 3 commercial data for quality assurance")
    add_bullet(doc, "Cost: ~$98,830/year (incremental)")

    doc.add_heading("Phase 5 — Steady State (Week 8+)", level=2)
    add_bullet(doc, "All 51 jurisdictions ingesting on their defined schedules")
    add_bullet(doc, "Monitor latency SLA dashboard daily")
    add_bullet(doc, "Evaluate South Carolina subscription ($54K) vs. commercial provider coverage")
    add_bullet(doc, "Evaluate PostgreSQL migration if data volumes or concurrent access demands increase")

    # ======================================================================
    # 12. RISK FACTORS & MITIGATIONS
    # ======================================================================
    doc.add_heading("12. Risk Factors & Mitigations", level=1)

    add_styled_table(
        doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            [
                "Commercial provider contract\nexceeds $300K estimate",
                "Medium",
                "Medium",
                "Budget buffer of $601K absorbs up to $500K.\nNegotiate multi-year discount."
            ],
            [
                "State changes bulk data\nformat without notice",
                "Medium",
                "Low",
                "Raw source data preserved; adapter can\nbe updated without data loss."
            ],
            [
                "Socrata API rate limiting\nor deprecation",
                "Low",
                "Low",
                "Only 2 states affected; fall back to\ncommercial provider for CT/CO."
            ],
            [
                "State bulk subscription\nprice increases",
                "Medium",
                "Low",
                "Budget buffer; can drop Tier 2 states\nand use commercial provider instead."
            ],
            [
                "Delaware authorized\nsearcher access",
                "High",
                "Low",
                "Use FCS (an existing Authorized Searcher)\nfor Delaware via commercial API."
            ],
            [
                "Data quality inconsistencies\nacross sources",
                "High",
                "Medium",
                "Normalizer layer + source_raw preservation\nfor audit. Cross-reference Tier 2 vs Tier 3."
            ],
            [
                "SQLite performance at scale\n(>5M records)",
                "Low",
                "Medium",
                "Designed for easy migration to PostgreSQL.\nSQLite WAL mode handles 1.7M/yr comfortably."
            ],
        ],
        col_widths=[3.5, 2.0, 1.5, 6.0],
    )

    # ======================================================================
    # 13. RECOMMENDATION
    # ======================================================================
    doc.add_heading("13. Recommendation", level=1)

    p = doc.add_paragraph(
        "We recommend proceeding with the hybrid 3-tier ingestion architecture as described "
        "in this document. The approach balances cost, data quality, and operational simplicity:"
    )

    doc.add_paragraph("")

    add_bullet(doc, "Begin Phase 1 immediately — test the pipeline against the free CT and CO APIs at zero cost and zero risk. This validates the entire ingestion, normalization, and storage pipeline end-to-end with real data.", bold_prefix="Start with free validation: ")

    doc.add_paragraph("")

    add_bullet(doc, "In parallel with Phase 1 testing, initiate contact with Baselayer, First Corporate Solutions, and LexisNexis to solicit formal quotes. This is the longest lead-time item and should start as early as possible.", bold_prefix="Begin commercial provider outreach: ")

    doc.add_paragraph("")

    add_bullet(doc, "Once the commercial provider is live, begin activating Tier 2 state subscriptions in priority order. Each Tier 2 source activated reduces dependence on the commercial provider and provides a higher-authority data source for cross-validation.", bold_prefix="Layer in state subscriptions: ")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Closing line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of Document")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    # -- Save --
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
